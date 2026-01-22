import streamlit as st
import pandas as pd
from datetime import date, datetime
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

# ---------------- CONFIG ----------------
st.set_page_config(page_title="Sistema Financeiro MEI", layout="centered")

BASE_PATH = Path(".")
DATA_PATH = BASE_PATH / "data" / "clientes"
OUTPUT_PATH = BASE_PATH / "outputs"
LOGS_PATH = BASE_PATH / "logs"
USUARIOS_FILE = BASE_PATH / "usuarios.csv"

for p in [DATA_PATH, OUTPUT_PATH, LOGS_PATH]:
    p.mkdir(parents=True, exist_ok=True)

# ---------------- LOGIN ----------------
if "logado" not in st.session_state:
    st.session_state.logado = False
    st.session_state.usuario = None

if not st.session_state.logado:
    st.title("🔐 Login")

    user = st.text_input("Usuário")
    senha = st.text_input("Senha", type="password")
    entrar = st.button("Entrar")

    if entrar:
        if USUARIOS_FILE.exists():
            df_users = pd.read_csv(USUARIOS_FILE)
            ok = df_users[
                (df_users["usuario"] == user) &
                (df_users["senha"] == senha)
            ]
            if not ok.empty:
                st.session_state.logado = True
                st.session_state.usuario = user
                st.experimental_rerun()
            else:
                st.error("Usuário ou senha inválidos")
        else:
            st.error("Arquivo usuarios.csv não encontrado")

    st.stop()

# ---------------- LOGOUT ----------------
colL, colR = st.columns([6,1])
with colR:
    if st.button("🚪 Sair"):
        st.session_state.logado = False
        st.session_state.usuario = None
        st.experimental_rerun()

cliente = st.session_state.usuario
st.title("💼 Sistema Financeiro MEI")
st.caption(f"Cliente: {cliente}")

# ---------------- FUNÇÕES ----------------
def registrar_log(acao):
    log_file = LOGS_PATH / "access_log.csv"
    log = pd.DataFrame([{
        "cliente": cliente,
        "acao": acao,
        "data_hora": datetime.now()
    }])
    if log_file.exists():
        df_log = pd.read_csv(log_file)
        df_log = pd.concat([df_log, log], ignore_index=True)
    else:
        df_log = log
    df_log.to_csv(log_file, index=False)

def salvar_lancamento(registro):
    arq = DATA_PATH / f"{cliente}.csv"
    if arq.exists():
        df = pd.read_csv(arq, parse_dates=["data"])
        df = pd.concat([df, registro], ignore_index=True)
    else:
        df = registro
    df.to_csv(arq, index=False)
    registrar_log("Lançamento adicionado")

def carregar_dados():
    arq = DATA_PATH / f"{cliente}.csv"
    if arq.exists():
        return pd.read_csv(arq, parse_dates=["data"])
    return pd.DataFrame(columns=["data","descricao","valor","tipo","categoria"])

def salvar_todos(df):
    arq = DATA_PATH / f"{cliente}.csv"
    df.to_csv(arq, index=False)

def gerar_pdf(df_mes, mes, ano, receitas, despesas, saldo):
    nome = f"relatorio_{cliente}_{mes}_{ano}.pdf"
    path = OUTPUT_PATH / nome

    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    e = []

    e.append(Paragraph("<b>Relatório Financeiro Mensal</b>", styles["Title"]))
    e.append(Paragraph(f"Cliente: {cliente}", styles["Normal"]))
    e.append(Paragraph(f"Mês/Ano: {mes}/{ano}", styles["Normal"]))
    e.append(Spacer(1, 10))
    e.append(Paragraph(f"Receitas: R$ {receitas:,.2f}", styles["Normal"]))
    e.append(Paragraph(f"Despesas: R$ {despesas:,.2f}", styles["Normal"]))
    e.append(Paragraph(f"Saldo: R$ {saldo:,.2f}", styles["Normal"]))
    e.append(Spacer(1, 10))

    tabela = [df_mes.columns.tolist()] + df_mes.astype(str).values.tolist()
    e.append(Table(tabela, repeatRows=1))
    doc.build(e)
    return path

def gerar_excel(df_mes, mes, ano):
    nome = f"relatorio_{cliente}_{mes}_{ano}.xlsx"
    path = OUTPUT_PATH / nome
    df_mes.to_excel(path, index=False)
    return path

def gerar_word(df_mes, mes, ano):
    nome = f"relatorio_{cliente}_{mes}_{ano}.docx"
    path = OUTPUT_PATH / nome
    doc = Document()
    doc.add_heading(f'Relatório Financeiro - {cliente}', 0)
    doc.add_paragraph(f"Mês/Ano: {mes}/{ano}")

    table = doc.add_table(rows=1, cols=len(df_mes.columns))
    for i, col in enumerate(df_mes.columns):
        table.rows[0].cells[i].text = col

    for _, row in df_mes.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df_mes.columns):
            cells[i].text = str(row[col])

    doc.save(path)
    return path

# ---------------- APP ----------------
df = carregar_dados()

st.subheader("📝 Novo lançamento")
with st.form("form_lanc"):
    c1, c2 = st.columns(2)
    with c1:
        data = st.date_input("Data", value=date.today())
        desc = st.text_input("Descrição")
    with c2:
        valor = st.number_input("Valor (R$)", step=0.01, format="%.2f")
        tipo = st.selectbox("Tipo", ["Despesa","Receita"])
        cat = st.selectbox("Categoria", ["Aluguel","Fornecedor","Combustível","Alimentação","Impostos","Outros"])
    ok = st.form_submit_button("Salvar")

if ok:
    novo = pd.DataFrame([{
        "data": data, "descricao": desc, "valor": valor, "tipo": tipo, "categoria": cat
    }])
    salvar_lancamento(novo)
    st.success("Salvo!")
    st.experimental_rerun()

if not df.empty:
    st.subheader("📊 Relatórios")

    df["mes"] = df["data"].dt.month
    df["ano"] = df["data"].dt.year

    mes = st.selectbox("Mês", sorted(df["mes"].unique()))
    ano = st.selectbox("Ano", sorted(df["ano"].unique()))

    df_mes = df[(df["mes"]==mes)&(df["ano"]==ano)]

    st.subheader("📋 Lançamentos (clique para excluir)")
    for i, row in df_mes.iterrows():
        col1, col2 = st.columns([8,1])
        with col1:
            st.write(f"{row['data'].date()} | {row['descricao']} | {row['tipo']} | R$ {row['valor']:.2f} | {row['categoria']}")
        with col2:
            if st.button("❌", key=f"del{i}"):
                df = df.drop(i)
                salvar_todos(df)
                st.experimental_rerun()

    receitas = df_mes[df_mes["tipo"]=="Receita"]["valor"].sum()
    despesas = df_mes[df_mes["tipo"]=="Despesa"]["valor"].sum()
    saldo = receitas - despesas

    c1,c2,c3 = st.columns(3)
    c1.metric("Receitas", f"R$ {receitas:,.2f}")
    c2.metric("Despesas", f"R$ {despesas:,.2f}")
    c3.metric("Saldo", f"R$ {saldo:,.2f}")

    st.subheader("⬇️ Downloads")
    pdf = gerar_pdf(df_mes, mes, ano, receitas, despesas, saldo)
    xls = gerar_excel(df_mes, mes, ano)
    docx = gerar_word(df_mes, mes, ano)

    st.download_button("PDF", open(pdf,"rb"), file_name=pdf.name)
    st.download_button("Excel", open(xls,"rb"), file_name=xls.name)
    st.download_button("Word", open(docx,"rb"), file_name=docx.name)

else:
    st.info("Nenhum lançamento ainda.")
